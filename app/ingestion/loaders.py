from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models import DocumentMetadata, DocumentType, NormalizedDocument


def _collapse_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        parts.append(extracted)
    return _collapse_whitespace("\n".join(parts))


def load_pdf(file_path: str | Path) -> str:
    path = Path(file_path)
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        parts.append(extracted)
    return _collapse_whitespace("\n".join(parts))


def load_docx(file_path: str | Path) -> str:
    path = Path(file_path)
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return _collapse_whitespace("\n".join(paragraphs))


def load_txt(file_path: str | Path) -> str:
    path = Path(file_path)
    raw = path.read_text(encoding="utf-8", errors="replace")
    return _collapse_whitespace(raw)


def load_html(file_path: str | Path) -> str:
    """Minimal HTML stripping without extra deps; stretch goal."""
    path = Path(file_path)
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"(?is)<script.*?>.*?</script>", " ", raw)
    raw = re.sub(r"(?is)<style.*?>.*?</style>", " ", raw)
    raw = re.sub(r"(?s)<[^>]+>", " ", raw)
    return _collapse_whitespace(raw)


def normalize_document(
    text: str,
    *,
    source: str,
    document_id: str,
    doc_type: DocumentType = "manual",
    species: str | None = None,
    topic: str | None = None,
    extra_metadata: dict[str, Any] | None = None,
) -> NormalizedDocument:
    meta = DocumentMetadata(type=doc_type, species=species, topic=topic)
    base: dict[str, Any] = meta.model_dump(exclude_none=True)
    if extra_metadata:
        for k, v in extra_metadata.items():
            if v is not None and k not in base:
                base[k] = v
    return NormalizedDocument(
        text=text,
        source=source,
        document_id=document_id,
        metadata=DocumentMetadata.model_validate(base),
    )
